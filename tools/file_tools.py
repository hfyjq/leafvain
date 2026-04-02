import os
import json
from pathlib import Path
from config import WORKSPACE_DIR, ALLOWED_EXTENSIONS, MAX_FILE_SIZE
import re
from typing import List, Dict, Any, Optional
from datetime import datetime
from core.tool_interface import ToolInterface
import config
from core.safety_checker import SafetyChecker
from core.api_client import APIClientFactory
from config import WORKSPACE_DIR
import sys
sys.path.append(str(Path(__file__).resolve().parent.parent))
# 在 file_tools.py 顶部添加
import os
import sys
from pathlib import Path

# 确保正确导入config
try:
    from config import WORKSPACE_DIR, ALLOWED_EXTENSIONS
except ImportError:
    # 备用方案：手动计算路径
    BASE_DIR = Path(__file__).resolve().parent.parent
    WORKSPACE_DIR = BASE_DIR / "workspace"
    ALLOWED_EXTENSIONS = {'.txt', '.md', '.py', '.json', '.csv', '.doc', '.docx', '.pdf'}
    print(f"⚠️ 备用配置: WORKSPACE_DIR={WORKSPACE_DIR}")


class FileTools(ToolInterface):
    """文件操作工具集"""



    def __init__(self):
        # 初始化API客户端
        super().__init__()
        self.metadata = {
            "name": "file_tools",
            "description": "文件操作工具集",
            "parameters": {
                "list_files": {
                    "directory_path": {"type": "string", "optional": True},
                    "recursive": {"type": "boolean", "optional": True, "default": False}
                },
                "read_file": {
                    "file_path": {"type": "string"},
                    "max_lines": {"type": "integer", "optional": True, "default": 1000}
                },
                "summarize_content": {
                    "content": {"type": "string"},
                    "max_length": {"type": "integer", "optional": True, "default": 500}
                }
            }
        }
        self.api_client = None

    def execute(self, **kwargs) -> Dict[str, Any]:
        action = kwargs.pop("action", None)
        if action == "list_files":
            return self.list_files(**kwargs)
        elif action == "read_file":
            return self.read_file(**kwargs)

    def _get_api_client(self):
        """获取API客户端（懒加载）"""
        if self.api_client is None:
            self.api_client = APIClientFactory.create_client("zhipu")
        return self.api_client

    def _simple_summary(self, content: str, max_length: int = 500) -> Dict[str, Any]:
        """简单的文本摘要（备用）"""
        if not content or len(content.strip()) == 0:
            return {
                "success": False,
                "error": "内容为空"
            }



        # 简单摘要算法：取前几行和非空行
        lines = [line.strip() for line in content.split('\n') if line.strip()]

        if len(lines) <= 5:
            summary = content
        else:
            # 取开头、中间和结尾部分
            first_lines = lines[:3]
            middle_lines = lines[len(lines) // 2 - 1:len(lines) // 2 + 2] if len(lines) > 6 else []
            last_lines = lines[-2:] if len(lines) > 5 else []

            summary_lines = first_lines
            if middle_lines:
                summary_lines.append("...")
                summary_lines.extend(middle_lines)
            if last_lines:
                if middle_lines:
                    summary_lines.append("...")
                summary_lines.extend(last_lines)

            summary = '\n'.join(summary_lines)

        # 限制长度
        if len(summary) > max_length:
            summary = summary[:max_length] + "..."

        return {
            "success": True,
            "summary": summary,
            "original_length": len(content),
            "summary_length": len(summary),
            "compression_ratio": f"{len(summary) / max(len(content), 1) * 100:.1f}%"
        }

    def summarize_large_file(self, file_path: str, max_chunk_size: int = 3000) -> Dict[str, Any]:
        """
        智能分块总结大文件
        """
        # 1. 读取文件内容
        read_result = self.read_file(file_path)
        if not read_result["success"]:
            return read_result

        content = read_result["content"]

        # 2. 智能分块（按章节结构）
        chunks = []
        current_chunk = ""
        for line in content.splitlines():
            # 检测章节标题（如"1."、"2.1"、"##"等格式）
            if re.match(r"^(\d+\.|\d+\.\d+|\#{1,3})\s", line.strip()):
                if current_chunk:  # 保存当前块
                    chunks.append(current_chunk)
                    current_chunk = ""
            current_chunk += line + "\n"

        if current_chunk:  # 添加最后一块
            chunks.append(current_chunk)

        # 3. 分块总结
        summaries = []
        for chunk in chunks:
            # 处理超大块（二次分块）
            if len(chunk) > max_chunk_size * 2:
                sub_chunks = [chunk[i:i + max_chunk_size] for i in range(0, len(chunk), max_chunk_size)]
                for sub in sub_chunks:
                    result = self.summarize_content(sub)
                    if result["success"]:
                        summaries.append(result["summary"])
            else:
                result = self.summarize_content(chunk)
                if result["success"]:
                    summaries.append(result["summary"])

        # 4. 合并总结
        full_summary = "\n\n".join(summaries)
        return {
            "success": True,
            "summary": full_summary,
            "original_length": len(content),
            "summary_length": len(full_summary),
            "chunks": len(chunks)
        }

    def summarize_content(self, content: str, max_length: int = 500) -> Dict[str, Any]:
        """
        使用大模型总结文本内容（增强修复版）
        """
        try:
            # 预处理：还原转义字符
            content = content.replace('\\n', '\n').replace('\\r', '\r').replace('\\t', '\t').replace('\\"', '"')

            if not content or len(content.strip()) == 0:
                return {
                    "success": False,
                    "error": "内容为空"
                }

            # 如果内容太短，使用简单总结
            if len(content) < 200:
                return self._simple_summary(content, max_length)

            # 调用大模型API进行总结
            client = self._get_api_client()

            # 构建更有效的提示词（添加系统消息）
            messages = [
                {
                    "role": "system",
                    "content": "你是一个专业的文档总结助手。请用中文简洁地总结以下内容，突出重点信息，保持专业性和准确性。"
                },
                {
                    "role": "user",
                    "content": f"请总结以下内容（不超过{max_length}字）：\n\n{content[:5000]}"  # 限制输入长度
                }
            ]

            print(f"🤔 使用大模型总结内容 ({len(content)} 字符)...")

            # 调用大模型API
            response = client.chat_completion(
                messages=messages,
                temperature=0.1,
                max_tokens=1000
            )

            # 关键修复：正确处理API响应对象
            if hasattr(response, 'choices') and response.choices:
                summary = response.choices[0].message.content.strip()
            else:
                # 如果响应结构异常，尝试直接提取内容
                summary = str(response).strip()
                if len(summary) > max_length:
                    summary = summary[:max_length] + "..."

            return {
                "success": True,
                "summary": summary,
                "original_length": len(content),
                "summary_length": len(summary),
                "compression_ratio": f"{len(summary) / max(len(content), 1) * 100:.1f}%"
            }
        except Exception as e:
            print(f"⚠️ 大模型总结失败: {str(e)}")
            return self._simple_summary(content, max_length)



    @staticmethod
    def normalize_extension(self, ext: str) -> str:
        """简化扩展名规范化"""
        if not ext:
            return ""

        # 转换为小写并去除空格
        ext = ext.strip().lower()

        # 确保以点开头
        if not ext.startswith('.'):
            ext = '.' + ext

        return ext

    def _human_readable_size(self, size_bytes: int) -> str:
        """转换字节数为易读格式"""
        if size_bytes == 0:
            return "0B"
        units = ("B", "KB", "MB", "GB")
        i = 0
        while size_bytes >= 1024 and i < len(units) - 1:
            size_bytes /= 1024
            i += 1
        return f"{size_bytes:.2f}{units[i]}"

    def list_files(self, directory_path: str = "", recursive: bool = False) -> Dict[str, Any]:
        try:
            # 处理特殊情况
            if not directory_path or directory_path.strip() == "" or "{workspace_path}" in directory_path:
                target_dir = WORKSPACE_DIR
            else:
                # 使用旧版路径验证
                is_valid, path_result = SafetyChecker.validate_path(directory_path, allow_directory=True)
                if not is_valid:
                    return {
                        "success": False,
                        "error": f"路径验证失败: {path_result}",
                        "files": []
                    }
                target_dir = path_result

            print(f"✅ 最终目标目录: {target_dir}")

            # 检查目录是否存在
            if not target_dir.exists() or not target_dir.is_dir():
                return {
                    "success": False,
                    "error": f"不是有效目录: {target_dir}",
                    "files": []
                }

            # 列出文件（恢复旧版逻辑）
            files = []
            dirs = []

            for item in target_dir.iterdir():
                try:
                    if item.is_file():
                        stat = item.stat()
                        files.append({
                            "name": item.name,
                            "path": str(item.relative_to(WORKSPACE_DIR)),
                            "size": stat.st_size,
                            "size_human": self._human_readable_size(stat.st_size),
                            "modified": stat.st_mtime,
                            "modified_date": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                            "extension": item.suffix.lower()  # 恢复旧版扩展名处理
                        })
                    elif item.is_dir() and recursive:
                        dirs.append({
                            "name": item.name,
                            "path": str(item.relative_to(WORKSPACE_DIR)),
                            "type": "directory"
                        })
                except (PermissionError, OSError) as e:
                    print(f"⚠️ 忽略无法访问的文件/目录: {item.name}")
                    continue

            # 按修改时间排序（最新优先）
            files.sort(key=lambda x: x["modified"], reverse=True)

            # 构建返回结果
            total_size = sum(f["size"] for f in files)
            total_size_human = self._human_readable_size(total_size)

            safe_target_dir = str(target_dir).replace('{', '{{').replace('}', '}}')

            # 构建安全格式化字符串
            formatted_result = """📂 目录: {}
            文件总数: {} 个
            总大小: {}""".format(
                safe_target_dir,
                len(files),
                total_size_human
            )

            return {
                "success": True,
                "files": files[:50],
                "directories": dirs[:20] if recursive else [],
                "count": len(files),
                "directory": str(target_dir.relative_to(WORKSPACE_DIR)),
                "total_size": total_size,
                "total_size_human": total_size_human,
                "formatted_result": formatted_result  # 修复后的格式化结果
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"列出文件时出错: {str(e)}",
                "files": []
            }

    def read_file(self, file_path: str, max_lines: int = 1000) -> Dict[str, Any]:
        """安全读取文件内容（增强大文件支持）"""
        try:
            # 验证路径
            is_valid, path_result = SafetyChecker.validate_path(file_path, allow_directory=False)
            if not is_valid:
                return {
                    "success": False,
                    "error": f"路径验证失败: {path_result}"
                }

            abs_path = path_result
            file_size = abs_path.stat().st_size
            is_large_file = file_size > config.MAX_CONTENT_LENGTH  # 10KB阈值

            # 检查文件操作
            is_safe, msg = SafetyChecker.check_file_operation(abs_path, "read")
            if not is_safe:
                return {
                    "success": False,
                    "error": f"文件安全检查失败: {msg}"
                }

            ext = abs_path.suffix.lower()
            if ext == '.pdf':
                return self._read_pdf(abs_path, max_lines, is_large_file)
            elif ext in ['.doc', '.docx']:
                return self._read_docx(abs_path, max_lines, is_large_file)
            else:
                # 文本文件处理
                return self._read_text_file(abs_path, max_lines, is_large_file)

        except Exception as e:
            return {
                "success": False,
                "error": f"读取文件时出错: {str(e)}"
            }

    def _read_text_file(self, abs_path: Path, max_lines: int, is_large_file: bool) -> Dict[str, Any]:
        """读取文本文件（增强大文件支持）"""
        try:
            try:
                with open(abs_path, 'r', encoding='utf-8', errors='ignore') as f:
                    lines = []
                    for i, line in enumerate(f):
                        if i >= max_lines:
                            lines.append(f"... [文件过长，已截断前{max_lines}行]")
                            break
                        lines.append(line.rstrip('\n'))
                    content = '\n'.join(lines)
            except UnicodeDecodeError:
                # 尝试其他编码
                with open(abs_path, 'r', encoding='gbk', errors='ignore') as f:
                    lines = []
                    for i, line in enumerate(f):
                        if i >= max_lines:
                            lines.append(f"... [文件过长，已截断前{max_lines}行]")
                            break
                        lines.append(line.rstrip('\n'))
                    content = '\n'.join(lines)

            # 获取文件信息
            stat = abs_path.stat()

            return {
                "success": True,
                "content": content,
                "file_name": abs_path.name,
                "file_path": str(abs_path.relative_to(WORKSPACE_DIR)),
                "file_size": len(content),
                "file_size_bytes": stat.st_size,
                "line_count": len(content.split('\n')),
                "modified": stat.st_mtime,
                "is_truncated": len(content.split('\n')) >= max_lines,
                "is_large_file": is_large_file  # 新增关键字段
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取文本文件时出错: {str(e)}"
            }

    def _read_pdf(self, file_path: Path, max_lines: int, is_large_file: bool) -> Dict[str, Any]:
        """读取PDF文件（文献优化版）"""
        try:
            from PyPDF2 import PdfReader
            content_lines = []

            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                total_pages = len(reader.pages)

                # 文献专用策略
                if "reference" in file_path.name.lower() or "literature" in file_path.name.lower():
                    # 优先读取：摘要(Abstract)、结论(Conclusion)、参考文献(References)
                    key_pages = []
                    for i in range(total_pages):
                        text = reader.pages[i].extract_text()
                        if re.search(r"abstract|摘要", text, re.IGNORECASE):
                            key_pages.append(i)
                        if re.search(r"conclusion|结论", text, re.IGNORECASE):
                            key_pages.append(i)
                        if re.search(r"references|参考文献", text, re.IGNORECASE):
                            key_pages.append(i)

                    # 确保至少包含首尾页
                    if 0 not in key_pages:
                        key_pages.insert(0, 0)
                    if total_pages - 1 not in key_pages:
                        key_pages.append(total_pages - 1)
                else:
                    # 通用策略（原有逻辑）
                    key_pages = self._select_key_pages(total_pages, min(total_pages, 10))

                # 读取关键页
                for page_num in key_pages:
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        lines = text.splitlines()
                        content_lines.extend(lines[:max_lines])

            content = "\n".join(content_lines[:max_lines])
            return {
                "success": True,
                "content": content,
                "file_type": "pdf",
                "key_pages": key_pages,
                "total_pages": total_pages
            }
        except Exception as e:
            return {"success": False, "error": f"PDF解析失败: {str(e)}"}

    def _select_key_pages(self, total_pages: int, max_pages: int) -> List[int]:
        """智能选择关键页面（首尾页+中间页）"""
        if total_pages <= max_pages:
            return list(range(total_pages))

        # 优先选择：开头3页 + 结尾3页 + 中间均匀采样
        key_pages = [0, 1, 2]  # 开头
        key_pages.extend([total_pages - 3, total_pages - 2, total_pages - 1])  # 结尾

        # 中间均匀采样
        step = max(1, (total_pages - 6) // (max_pages - 6))
        for i in range(3, total_pages - 3, step):
            if len(key_pages) >= max_pages:
                break
            key_pages.append(i)

        return sorted(set(key_pages))[:max_pages]

    def _read_docx(self, file_path: Path, max_lines: int, is_large_file: bool) -> Dict[str, Any]:
        """读取DOCX文件（智能分段策略）"""
        try:
            from docx import Document
            doc = Document(file_path)
            content_lines = []
            total_paragraphs = len(doc.paragraphs)

            # 智能段落选择策略
            if is_large_file and total_paragraphs > 20:
                # 关键段落选择：开头10% + 核心章节 + 结尾10%
                start_index = int(total_paragraphs * 0.1)
                end_index = total_paragraphs - int(total_paragraphs * 0.1)

                # 添加开头部分
                for i in range(0, min(5, start_index)):  # 至少前5段
                    content_lines.append(doc.paragraphs[i].text)

                # 添加核心章节（检测标题格式）
                for i in range(start_index, end_index):
                    text = doc.paragraphs[i].text.strip()
                    # 检测章节标题（如"1."、"2.1"等）
                    if re.match(r"^(\d+\.|\d+\.\d+|\#{1,3})\s", text):
                        content_lines.append("\n" + text)  # 添加标题
                        # 添加标题后3段内容
                        for j in range(i + 1, min(i + 4, end_index)):
                            content_lines.append(doc.paragraphs[j].text)

                # 添加结尾部分
                for i in range(max(end_index, total_paragraphs - 3), total_paragraphs):
                    content_lines.append(doc.paragraphs[i].text)
            else:
                # 小文件全量读取
                for para in doc.paragraphs:
                    content_lines.append(para.text)

            content = "\n".join(content_lines[:max_lines])
            return {
                "success": True,
                "content": content,
                "file_type": "docx",
                "is_large_file": is_large_file,
                "strategy": "smart_sampling" if is_large_file else "full"
            }
        except Exception as e:
            return {"success": False, "error": f"DOCX解析失败: {str(e)}"}

    @staticmethod
    def _human_readable_size(size_bytes: int) -> str:
        """将字节数转换为人类可读的格式"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f}PB"
